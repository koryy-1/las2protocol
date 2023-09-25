from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from calc_types import ParametersForReporting


def make_docx(params_for_reporting: ParametersForReporting, MEM_FILES, picture_size):
    (MF_NEAR_PROBE, MF_FAR_PROBE, MF_FAR_ON_NEAR_PROBE, MF_MT) = MEM_FILES

    document = Document()

    # styles
    style = document.styles['Normal']
    style.font.size = Pt(12)

    p_header = document.add_paragraph()
    p_header.add_run('Протокол').bold = True
    p_header_fmt = p_header.paragraph_format
    p_header_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_header2 = document.add_paragraph('температурных испытаний прибора')
    p_header2_fmt = p_header2.paragraph_format
    p_header2_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_paragraph(f'Прибор:                               №: {params_for_reporting.serial_number}', style='List Number')
    document.add_paragraph(f'Канал:                                   {params_for_reporting.instrument_name}', style='List Number')
    document.add_paragraph(f'Дата испытаний:                 {params_for_reporting.date}', style='List Number')
    document.add_paragraph(f'Пороги: {params_for_reporting.near_probe_title} - {params_for_reporting.near_probe_threshold} мВ, {params_for_reporting.far_probe_title} - {params_for_reporting.far_probe_threshold} мВ. ', style='List Number')

    document.add_paragraph(params_for_reporting.near_probe_title, style='List Number')
    document.add_picture(MF_NEAR_PROBE, width=Inches(picture_size))

    document.add_paragraph(params_for_reporting.far_probe_title, style='List Number')
    document.add_picture(MF_FAR_PROBE, width=Inches(picture_size))

    document.add_paragraph(f'{params_for_reporting.far_probe_title}/{params_for_reporting.near_probe_title}', style='List Number')
    document.add_picture(MF_FAR_ON_NEAR_PROBE, width=Inches(picture_size))

    document.add_paragraph('TEMPER', style='List Number')
    document.add_picture(MF_MT, width=Inches(picture_size))


    # for heating
    if params_for_reporting.heating_table is not None:
        document.add_paragraph('Результаты при нагреве', style='List Number')

        cooling_table = document.add_table(rows=1, cols=5)
        hdr_cells = cooling_table.rows[0].cells
        hdr_cells[0].text = 'п/п'
        hdr_cells[0].width = Inches(0.4)
        hdr_cells[1].text = 'Формула'
        hdr_cells[1].width = Inches(1.5)
        hdr_cells[2].text = params_for_reporting.near_probe_title
        hdr_cells[3].text = params_for_reporting.far_probe_title
        hdr_cells[4].text = f'{params_for_reporting.far_probe_title}/{params_for_reporting.near_probe_title}'
        # print(records)
        for num, formula, NEAR_PROBE, FAR_PROBE, FAR_ON_NEAR_PROBE in params_for_reporting.heating_table:
            row_cells = cooling_table.add_row().cells
            row_cells[0].text = str(num)
            row_cells[0].width = Inches(0.4)
            row_cells[1].text = formula
            row_cells[1].width = Inches(1.5)
            row_cells[2].text = str(NEAR_PROBE)
            row_cells[3].text = str(FAR_PROBE)
            row_cells[4].text = str(FAR_ON_NEAR_PROBE)

        cooling_table.style = 'Table Grid'


    # for cooling
    if params_for_reporting.cooling_table is not None:
        document.add_paragraph('Результаты при охлаждении', style='List Number')

        cooling_table = document.add_table(rows=1, cols=5)
        hdr_cells = cooling_table.rows[0].cells
        hdr_cells[0].text = 'п/п'
        hdr_cells[0].width = Inches(0.4)
        hdr_cells[1].text = 'Формула'
        hdr_cells[1].width = Inches(1.5)
        hdr_cells[2].text = params_for_reporting.near_probe_title
        hdr_cells[3].text = params_for_reporting.far_probe_title
        hdr_cells[4].text = f'{params_for_reporting.far_probe_title}/{params_for_reporting.near_probe_title}'
        # print(records)
        for num, formula, NEAR_PROBE, FAR_PROBE, FAR_ON_NEAR_PROBE in params_for_reporting.cooling_table:
            row_cells = cooling_table.add_row().cells
            row_cells[0].text = str(num)
            row_cells[0].width = Inches(0.4)
            row_cells[1].text = formula
            row_cells[1].width = Inches(1.5)
            row_cells[2].text = str(NEAR_PROBE)
            row_cells[3].text = str(FAR_PROBE)
            row_cells[4].text = str(FAR_ON_NEAR_PROBE)

        cooling_table.style = 'Table Grid'
    
    
    document.add_paragraph()

    temp_ranges = get_temp_ranges(params_for_reporting)

    conc_p = document.add_paragraph(style='List Number')
    conc_p.add_run('Выводы').bold = True
    document.add_paragraph(
        f'Температурный уход сигналов {params_for_reporting.near_probe_title},  {params_for_reporting.far_probe_title}  и  {params_for_reporting.far_probe_title}/{params_for_reporting.near_probe_title}  в диапазоне температур {temp_ranges} {params_for_reporting.conclusion} 5%.'
    )

    document.add_paragraph()

    document.add_paragraph('Термоиспытания провел:')
    document.add_paragraph('Представитель ОТК:')

    return document


def get_temp_ranges(params_for_reporting: ParametersForReporting):
    heating_range = f'от {int(params_for_reporting.temp_min_left)} до {int(params_for_reporting.temp_max)} градусов'
    cooling_range = f'от {int(params_for_reporting.temp_max)} до {int(params_for_reporting.temp_min_right)} градусов'
    return f'{heating_range if params_for_reporting.heating_table else ""}{" и " if params_for_reporting.heating_table and params_for_reporting.cooling_table else ""}{cooling_range if params_for_reporting.cooling_table else ""}'
